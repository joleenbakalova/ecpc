#!/usr/bin/env python3
"""Convert an ECPC Markdown deliverable to a styled, editable Word (.docx).

Local-only (python-docx + stdlib zipfile); no external API. Handles the subset
of Markdown used in the ECPC briefs: #/##/### headings, > blockquote, - bullets,
1. numbered lists, --- rules, inline **bold** / *italic* / `code`, and
Markdown-style footnotes: an inline [^1] marker plus a [^1]: citation line
(placed anywhere, conventionally at the end). Footnotes become real Word
footnotes rendered at the bottom of the page.

Usage:  python md_to_docx.py INPUT.md OUTPUT.docx
"""
import re, sys, os, zipfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN_D = RGBColor(0x1B, 0x4B, 0x3A)   # dark green (titles)
GREEN   = RGBColor(0x2E, 0x7D, 0x64)   # ECPC green (headings)
GREY    = RGBColor(0x55, 0x55, 0x55)

INLINE = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*|\[\^\w+\])')
FNDEF  = re.compile(r'^\[\^(\w+)\]:\s*(.*)$')

FNMAP = {}   # footnote label -> integer id (set per build)

def add_footnote_ref(p, fid):
    run = p.add_run()
    rpr = run._r.get_or_add_rPr()
    v = OxmlElement('w:vertAlign'); v.set(qn('w:val'), 'superscript'); rpr.append(v)
    ref = OxmlElement('w:footnoteReference'); ref.set(qn('w:id'), str(fid))
    run._r.append(ref)

def add_runs(p, text):
    """Add text to paragraph p, honoring **bold**, *italic*, `code`, [^fn]."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = p.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        elif tok.startswith('*') and tok.endswith('*'):
            r = p.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith('[^') and tok.endswith(']'):
            fid = FNMAP.get(tok[2:-1])
            if fid:
                add_footnote_ref(p, fid)
        else:
            p.add_run(tok)

def bottom_border(p):
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
    for k, v in (('w:val','single'),('w:sz','6'),('w:space','1'),('w:color','C9C9C9')):
        b.set(qn(k), v)
    pbdr.append(b); pPr.append(pbdr)

def left_border(p):
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); l = OxmlElement('w:left')
    for k, v in (('w:val','single'),('w:sz','18'),('w:space','12'),('w:color','2E7D64')):
        l.set(qn(k), v)
    pbdr.append(l); pPr.append(pbdr)

def style_font(doc, name, size, color, bold=True, italic=False, before=0, after=4):
    st = doc.styles[name]
    st.font.name = 'Calibri'; st.font.size = Pt(size); st.font.bold = bold
    st.font.italic = italic; st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

def esc(s):
    return s.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')

def inject_footnotes(docx_path, notes):
    """notes: list of (id:int, plain-text citation). Adds a real footnotes part."""
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    parts = ['<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
             f'<w:footnotes xmlns:w="{W}">',
             '<w:footnote w:type="separator" w:id="-1"><w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p></w:footnote>',
             '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>']
    for fid, text in notes:
        parts.append(
            f'<w:footnote w:id="{fid}"><w:p><w:pPr><w:spacing w:after="40"/></w:pPr>'
            f'<w:r><w:rPr><w:vertAlign w:val="superscript"/><w:sz w:val="18"/></w:rPr><w:footnoteRef/></w:r>'
            f'<w:r><w:rPr><w:sz w:val="18"/></w:rPr><w:t xml:space="preserve"> {esc(text)}</w:t></w:r>'
            f'</w:p></w:footnote>')
    parts.append('</w:footnotes>')
    footnotes_xml = ''.join(parts).encode('utf-8')

    tmp = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin:
        names = zin.namelist()
        rels = zin.read('word/_rels/document.xml.rels').decode('utf-8')
        ct = zin.read('[Content_Types].xml').decode('utf-8')
        ids = [int(x) for x in re.findall(r'Id="rId(\d+)"', rels)]
        newid = 'rId%d' % ((max(ids) + 1) if ids else 1)
        rels = rels.replace('</Relationships>',
            f'<Relationship Id="{newid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/></Relationships>')
        ct = ct.replace('</Types>',
            '<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/></Types>')
        with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
            for n in names:
                data = zin.read(n)
                if n == 'word/_rels/document.xml.rels': data = rels.encode('utf-8')
                elif n == '[Content_Types].xml': data = ct.encode('utf-8')
                zout.writestr(n, data)
            zout.writestr('word/footnotes.xml', footnotes_xml)
    os.replace(tmp, docx_path)

def build(md_path, out_path):
    raw_lines = open(md_path, encoding='utf-8').read().split('\n')

    # first pass: pull footnote definitions out of the body, number by order
    lines, notes = [], []
    FNMAP.clear()
    for ln in raw_lines:
        m = FNDEF.match(ln.strip())
        if m:
            fid = len(FNMAP) + 1
            FNMAP[m.group(1)] = fid
            notes.append((fid, m.group(2).strip()))
        else:
            lines.append(ln)

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'; normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.12
    style_font(doc, 'Title',     23, GREEN_D, after=2)
    style_font(doc, 'Subtitle',  14, GREEN,  after=10)
    style_font(doc, 'Heading 1', 15, GREEN_D, before=14, after=4)
    style_font(doc, 'Heading 2', 12, GREEN,  before=8,  after=2)

    seen_subtitle = False
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith('# '):
            add_runs(doc.add_paragraph(style='Title'), line[2:].strip())
        elif line.startswith('## '):
            if not seen_subtitle:
                seen_subtitle = True
                add_runs(doc.add_paragraph(style='Subtitle'), line[3:].strip())
            else:
                add_runs(doc.add_paragraph(style='Heading 1'), line[3:].strip())
        elif line.startswith('### '):
            add_runs(doc.add_paragraph(style='Heading 2'), line[4:].strip())
        elif line.startswith('> '):
            p = doc.add_paragraph(); left_border(p)
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
            add_runs(p, line[2:].strip())
        elif line.strip() == '---':
            p = doc.add_paragraph(); bottom_border(p); p.paragraph_format.space_after = Pt(2)
        elif re.match(r'^\d+\.\s', line):
            add_runs(doc.add_paragraph(style='List Number'), re.sub(r'^\d+\.\s', '', line))
        elif line.startswith('- '):
            add_runs(doc.add_paragraph(style='List Bullet'), line[2:].strip())
        else:
            p = doc.add_paragraph()
            if line.strip().startswith('*') and line.strip().endswith('*') and '**' not in line:
                r = p.add_run(line.strip()[1:-1]); r.italic = True
                r.font.color.rgb = GREY; r.font.size = Pt(10)
            else:
                add_runs(p, line.strip())

    doc.save(out_path)
    if notes:
        inject_footnotes(out_path, notes)
    print('Wrote', out_path, f'({len(notes)} footnotes)')

if __name__ == '__main__':
    build(sys.argv[1], sys.argv[2])
